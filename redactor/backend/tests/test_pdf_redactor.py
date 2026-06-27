from pathlib import Path
from types import SimpleNamespace
import sys
import tempfile

import pymupdf as fitz

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from pdf_redactor import utils
from pdf_redactor.detectors import (
    BICDetector,
    CreditCardDetector,
    DateDetector,
    EmailDetector,
    IBANDetector,
    IFSCDetector,
    PANDetector,
    SSNDetector,
    TimestampDetector,
    UPIVPADetector,
    registry,
)
from pdf_redactor.detectors.base import Detection
from pdf_redactor.engine import RedactionEngine, RedactionOptions
from pdf_redactor.processors.pdf import PDFProcessor
from pdf_redactor.cli import default_output_path, resolve_single_file_output_path, validate_output_flag


def make_options(**overrides):
    defaults = {
        "detectors": [],
        "ml": False,
        "ml_detectors": ["name", "address"],
        "aggressive": False,
        "text": None,
        "color": "black",
        "color_hex": None,
        "text_color": "white",
        "text_color_hex": None,
        "preview": False,
        "geographic_code": None,
        "custom_masks": [],
    }
    defaults.update(overrides)
    return RedactionOptions(**defaults)


def test_find_timestamp_returns_full_matches():
    detector = TimestampDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["Start 12:34 and end 9:05."])
    texts = [d.text for d in detections]
    assert "12:34" in texts
    assert "9:05" in texts


def test_find_bics_rejects_common_stop_words():
    detector = BICDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["DEPOSITS and WITHDRAWALS are not BICs."])
    assert not detections


def test_find_bics_matches_eight_and_eleven_character_codes():
    detector = BICDetector()
    doc = fitz.open()
    doc.new_page()
    # DEUTDEFF is a real Deutsche Bank Frankfurt BIC.
    detections = detector.detect(doc, ["Codes: DEUTDEFF and DEUTDEFF500"])
    texts = [d.text for d in detections]
    assert "DEUTDEFF" in texts
    assert "DEUTDEFF500" in texts


def test_find_custom_mask_escapes_regex_characters():
    from pdf_redactor.detectors.text import MaskDetector

    pattern = MaskDetector._build_pattern("C++")
    assert pattern.findall("C++ is allowed. C+ is not.") == ["C++"]


def test_default_output_path_stays_next_to_input():
    assert default_output_path("/tmp/example/input.pdf") == "/tmp/example/input_redacted.pdf"


def test_resolve_single_file_output_path_accepts_directory():
    output = resolve_single_file_output_path("/tmp/example/input.pdf", "/tmp/redacted")
    # Accept either forward or backward slash because we are on Windows.
    assert output.replace("\\", "/") == "/tmp/redacted/input_redacted.pdf"


def test_validate_output_flag_accepts_directory_for_single_file(tmp_path):
    args = SimpleNamespace(output=str(tmp_path / "redacted"))
    validate_output_flag(args.output, input_is_dir=False)
    assert (tmp_path / "redacted").is_dir()


def test_iban_checksum_rejects_random_tokens():
    # GB82WEST12345698765432 is a valid IBAN.
    assert utils.iban_checksum("GB82 WEST 1234 5698 7654 32")
    # Random token should fail.
    assert not utils.iban_checksum("XX00 0000 0000 0000")


def test_date_detector_rejects_impossible_dates():
    detector = DateDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["Dates: 41/1-10 and 37-1-18 and 05-08-2025"])
    texts = [d.text for d in detections]
    assert "05-08-2025" in texts
    assert "41/1-10" not in texts
    assert "37-1-18" not in texts


def test_email_detector_rejects_url_fragments():
    detector = EmailDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["Contact: user@example.com or path/to/user@example.com/file"])
    texts = [d.text for d in detections]
    assert "user@example.com" in texts
    assert "path/to/user@example.com/file" not in texts


def test_ifsc_detector():
    detector = IFSCDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["IFSC: ICIC0001234 and INVALID0000"])
    texts = [d.text for d in detections]
    assert "ICIC0001234" in texts
    assert "INVALID0000" not in texts


def test_pan_detector():
    detector = PANDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["PAN: ABCDE1234F"])
    assert any(d.text == "ABCDE1234F" for d in detections)


def test_upi_vpa_detector():
    detector = UPIVPADetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["UPI: user@okaxis paid user@unknown"])
    texts = [d.text for d in detections]
    assert "user@okaxis" in texts
    assert "user@unknown" not in texts


def test_engine_applies_redaction():
    doc = fitz.open()
    page = doc.new_page(width=200, height=200)
    page.insert_text((50, 50), " confidential ", fontsize=12)

    engine = RedactionEngine()
    options = make_options(detectors=["mask"], custom_masks=["confidential"])
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        output_path = Path(tmpdir) / "output.pdf"
        doc.save(input_path)
        doc.close()
        report = engine.process(str(input_path), str(output_path), options)

        assert output_path.exists()
        assert any(d.detector == "mask" for d in report.detections)
        redacted = fitz.open(output_path)
        text = redacted.load_page(0).get_text("text")
        redacted.close()
        assert "confidential" not in text.lower()


def test_registry_contains_all_detectors():
    names = registry.names()
    for expected in [
        "phone", "email", "iban", "bic", "timestamp", "date",
        "link", "barcode", "qrcode", "account", "ifsc", "micr",
        "upi", "pan", "aadhaar", "balance", "receiver", "name", "address",
        "ssn", "passport", "creditcard", "secret", "customer_id",
    ]:
        assert expected in names


def test_excel_processor_redacts_cell_values():
    from openpyxl import Workbook, load_workbook
    from pdf_redactor.processors import registry as proc_registry

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Contact: user@example.com"
    ws["A2"] = "SSN: 078-05-1120"
    ws["A3"] = "Normal text"

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.xlsx"
        output_path = Path(tmpdir) / "output.xlsx"
        wb.save(input_path)

        engine = RedactionEngine()
        options = make_options(detectors=["email", "ssn"])
        report = engine.process(str(input_path), str(output_path), options)

        out_wb = load_workbook(output_path)
        assert "user@example.com" not in str(out_wb.active["A1"].value)
        assert "078-05-1120" not in str(out_wb.active["A2"].value)
        assert out_wb.active["A3"].value == "Normal text"
        assert report.summary.get("email") or report.summary.get("ssn")


def test_word_processor_redacts_paragraph_text():
    from docx import Document
    from pdf_redactor.processors import registry as proc_registry

    doc = Document()
    doc.add_paragraph("Contact: user@example.com")
    doc.add_paragraph("SSN: 078-05-1120")
    doc.add_paragraph("Normal text")

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.docx"
        output_path = Path(tmpdir) / "output.docx"
        doc.save(input_path)

        engine = RedactionEngine()
        options = make_options(detectors=["email", "ssn"])
        report = engine.process(str(input_path), str(output_path), options)

        out_doc = Document(output_path)
        full_text = "\n".join(p.text for p in out_doc.paragraphs)
        assert "user@example.com" not in full_text
        assert "078-05-1120" not in full_text
        assert "Normal text" in full_text
        assert report.summary.get("email") or report.summary.get("ssn")


def test_credit_card_detector_with_luhn_check():
    detector = CreditCardDetector()
    doc = fitz.open()
    doc.new_page()
    # 4532015112830366 is a valid test Visa number.
    detections = detector.detect(doc, ["Card: 4532015112830366 and 1234567890123456"])
    texts = [d.text for d in detections]
    assert "4532015112830366" in texts
    assert "1234567890123456" not in texts


def test_ssn_detector():
    detector = SSNDetector()
    doc = fitz.open()
    doc.new_page()
    detections = detector.detect(doc, ["SSN: 078-05-1120"])
    assert any(d.text == "078-05-1120" for d in detections)


def test_detection_dataclass_requires_valid_confidence():
    try:
        Detection(page=0, text="x", detector="test", confidence="invalid")
        assert False, "Expected ValueError"
    except ValueError:
        pass




def test_pdf_processor_preview_returns_detections_without_redacting():
    doc = fitz.open()
    page = doc.new_page(width=200, height=200)
    page.insert_text((50, 50), " confidential ", fontsize=12)

    engine = RedactionEngine()
    options = make_options(detectors=["mask"], custom_masks=["confidential"], preview=True)
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        output_path = Path(tmpdir) / "output.pdf"
        doc.save(input_path)
        doc.close()
        report = engine.process(str(input_path), str(output_path), options)

        assert any(d.detector == "mask" for d in report.detections)
        # Preview should not create or modify the output file.
        assert not output_path.exists()
        original = fitz.open(input_path)
        text = original.load_page(0).get_text("text")
        original.close()
        assert "confidential" in text.lower()


def test_pdf_processor_apply_selections_redacts_only_selected():
    doc = fitz.open()
    page = doc.new_page(width=200, height=200)
    page.insert_text((50, 50), " alpha ", fontsize=12)
    page.insert_text((50, 100), " beta ", fontsize=12)

    engine = RedactionEngine()
    detect_options = make_options(detectors=["mask"], custom_masks=["alpha", "beta"])
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        output_path = Path(tmpdir) / "output.pdf"
        doc.save(input_path)
        doc.close()

        # Detect both words.
        preview_options = make_options(detectors=["mask"], custom_masks=["alpha", "beta"], preview=True)
        preview_report = engine.process(str(input_path), str(output_path), preview_options)
        assert len(preview_report.detections) == 2

        # Apply only the 'alpha' selection.
        alpha = next(d for d in preview_report.detections if "alpha" in d.text)
        processor = PDFProcessor(engine, detect_options)
        processor.apply_selections(
            str(input_path),
            str(output_path),
            [
                {
                    "page": alpha.page + 1,
                    "rects": [tuple(alpha.rects[0])],
                    "detector": alpha.detector,
                    "text": alpha.text,
                    "confidence": alpha.confidence,
                }
            ],
        )

        redacted = fitz.open(output_path)
        text = redacted.load_page(0).get_text("text")
        redacted.close()
        assert "alpha" not in text.lower()
        assert "beta" in text.lower()


def test_api_preview_and_apply():
    from fastapi.testclient import TestClient

    from api.main import app

    doc = fitz.open()
    page = doc.new_page(width=200, height=200)
    page.insert_text((50, 50), " user@example.com ", fontsize=12)
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        doc.save(input_path)
        doc.close()

        with TestClient(app) as client:
            with open(input_path, "rb") as f:
                preview_resp = client.post(
                    "/preview",
                    files={"file": ("input.pdf", f, "application/pdf")},
                    data={"options": '{"detectors":["email"]}'},
                )
            assert preview_resp.status_code == 200
            payload = preview_resp.json()
            assert "job_id" in payload
            assert "original_pdf_base64" in payload
            assert payload["report"]["pages"] == 1
            detections = payload["report"]["detections"]
            assert any(d["detector"] == "email" for d in detections)

            job_id = payload["job_id"]
            selection = detections[0]
            apply_resp = client.post(
                f"/apply/{job_id}",
                json={
                    "selections": [selection],
                    "options": {"detectors": ["email"], "color": "black"},
                },
            )
            assert apply_resp.status_code == 200
            apply_payload = apply_resp.json()
            assert "pdf_base64" in apply_payload
            assert "report" in apply_payload

            # Verify the redacted PDF no longer contains the email.
            import base64

            redacted_bytes = base64.b64decode(apply_payload["pdf_base64"])
            redacted_path = Path(tmpdir) / "redacted.pdf"
            redacted_path.write_bytes(redacted_bytes)
            redacted = fitz.open(redacted_path)
            text = redacted.load_page(0).get_text("text")
            redacted.close()
            assert "user@example.com" not in text
