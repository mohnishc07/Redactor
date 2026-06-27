"""Word processor using python-docx."""

from __future__ import annotations

import os
import re

from pdf_redactor.detectors.base import Detection

from .base import DocumentProcessor, registry


@registry.register
class WordProcessor(DocumentProcessor):
    extensions = ("docx",)

    def __init__(self, engine, options):
        super().__init__(options)
        self.engine = engine

    def process(self, input_path: str, output_path: str):
        from docx import Document

        from pdf_redactor.engine import RedactionReport

        doc = Document(input_path)
        try:
            text_pages = [self._extract_text(doc)]
            detections = self._detect(text_pages)
            self._apply(doc, detections)

            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            doc.save(output_path)

            report = RedactionReport(pages=1, detections=detections)
            report.summary = self._summarize(detections)
            return report
        finally:
            pass  # python-docx Document has no close method

    def _extract_text(self, doc) -> str:
        parts = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)

    def _detect(self, text_pages: list[str]) -> list[Detection]:
        from pdf_redactor.engine import build_detectors

        detectors = build_detectors(self.engine.registry, self.options)
        all_detections: list[Detection] = []
        dummy = _DummyDocument()
        for detector in detectors:
            for text in text_pages:
                all_detections.extend(detector.detect(dummy, [text]))

        if self.options.ml and self.engine.ml_registry is not None:
            for text in text_pages:
                all_detections.extend(
                    self.engine.ml_registry.detect(dummy, [text], self.options.ml_detectors)
                )

        if self.options.aggressive:
            return all_detections
        return [d for d in all_detections if d.confidence in {"high", "medium"}]

    def _apply(self, doc, detections: list[Detection]) -> None:
        redaction_text = self.options.text or "[REDACTED]"
        patterns = [re.escape(d.text) for d in detections]
        if not patterns:
            return
        combined = re.compile("(" + "|".join(patterns) + ")")

        for para in doc.paragraphs:
            self._redact_paragraph(para, combined, redaction_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._redact_paragraph(para, combined, redaction_text)

    def _redact_paragraph(self, para, combined, redaction_text: str) -> None:
        full_text = para.text
        new_text, count = combined.subn(redaction_text, full_text)
        if not count:
            return
        # Clear runs and set redacted text in the first run.
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = new_text
            else:
                run.text = ""

    def _summarize(self, detections: list[Detection]) -> dict[str, dict[str, int]]:
        summary: dict[str, dict[str, int]] = {}
        for d in detections:
            summary.setdefault(d.detector, {"high": 0, "medium": 0, "low": 0})
            summary[d.detector][d.confidence] += 1
        return summary


class _DummyDocument:
    def load_page(self, page_num: int):
        return _DummyPage()

    def __len__(self):
        return 1


class _DummyPage:
    def search_for(self, text: str):
        return []
